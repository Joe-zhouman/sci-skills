#!/usr/bin/env python3
"""Generate a standalone Highlights DOCX (Elsevier format).

Usage:
  python3 generate-highlights.py \\
    --highlight "Core finding one" \\
    --highlight "Core finding two" \\
    --highlight "Core finding three" \\
    --out highlights.docx

Requires: python-docx (pip install python-docx)
"""

import argparse

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    exit("Please install python-docx: pip install python-docx")


def set_paragraph(doc, text, bold=False, size=11, space_after=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def main():
    p = argparse.ArgumentParser(description="Generate a standalone Highlights DOCX")
    p.add_argument("--highlight", action="append", required=True,
                   help="One highlight bullet (repeat for multiple)")
    p.add_argument("--out", required=True, help="Output .docx path")

    args = p.parse_args()

    doc = Document()
    for section in doc.sections:
        from docx.shared import Inches
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    set_paragraph(doc, "Highlights", bold=True, size=14, space_after=12)

    for item in args.highlight:
        item = item.strip()
        if not item:
            continue
        if len(item) > 85:
            print(f"WARNING: '{item[:60]}...' is {len(item)} chars (max 85)",
                  file=__import__('sys').stderr)
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    out = args.out
    if not out.endswith(".docx"):
        out += ".docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
