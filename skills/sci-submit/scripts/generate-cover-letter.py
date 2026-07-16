#!/usr/bin/env python3
"""Generate a DOCX cover letter — first submission or revision.

Usage:
  # First submission
  python3 generate-cover-letter.py \\
    --type first \\
    --journal "Journal Name" \\
    --title "Full Manuscript Title Goes Here" \\
    --background "Two recent developments have converged..." \\
    --findings "We found that..." \\
    --journal-fit "Journal Name has published key studies at..." \\
    --audience "This work should interest researchers in..." \\
    --stats "~X,000 words, N figures, N tables" \\
    --declarations "This manuscript has not been published elsewhere..." \\
    --authors "First Author, Second Author, Corresponding Author*" \\
    --corresponding "email@university.edu (C. Author)" \\
    --dept "Department of XXXX" \\
    --institution "University of XXXX" \\
    --address "Address Line, City, Postal Code, Country" \\
    --out cover-letter.docx

  # Revision
  python3 generate-cover-letter.py \\
    --type revision \\
    --journal "Journal Name" \\
    --title "Full Manuscript Title Goes Here" \\
    --msid "JRNL-2026-XXXXX" \\
    --changes "Reformatted with journal template.\\nMinor language corrections." \\
    --stats "~X,000 words, N figures, N tables" \\
    --authors "First Author, Second Author, Corresponding Author*" \\
    --corresponding "email@university.edu (C. Author)" \\
    --dept "Department of XXXX" \\
    --institution "University of XXXX" \\
    --address "Address Line, City, Postal Code, Country" \\
    --out cover-letter-revision.docx

Requires: python-docx (pip install python-docx)
"""

import argparse
import os
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    exit("Please install python-docx: pip install python-docx")


def set_paragraph(doc, text, bold=False, size=11, space_after=0, space_before=0, alignment=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if alignment is not None:
        pf.alignment = alignment
    return p


def set_multiline(doc, text, size=11, space_after=6):
    for line in text.strip().split("\\n"):
        set_paragraph(doc, line.strip(), size=size, space_after=space_after)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(" " * 80)
    run.font.size = Pt(1)
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "12",
        qn("w:space"): "1",
        qn("w:color"): "000000",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_header(doc, args):
    set_paragraph(doc, args.dept, bold=True, size=12, space_after=0)
    set_paragraph(doc, args.institution, size=9, space_after=0)
    set_paragraph(doc, args.address, size=9, space_after=0, italic=True)
    add_horizontal_line(doc)
    set_paragraph(doc, date.today().strftime("%B %d, %Y"), size=11, space_after=12)


def build_first(doc, args):
    build_header(doc, args)
    set_paragraph(doc, "Dear Editor,", size=11, space_after=6)

    opening = (f'We wish to submit our manuscript entitled '
               f'"{args.title}" '
               f'for consideration for publication in {args.journal}.')
    set_paragraph(doc, opening, size=11, space_after=12)

    if args.background:
        set_paragraph(doc, args.background, size=11, space_after=12)
    if args.findings:
        set_paragraph(doc, args.findings, size=11, space_after=12)
    if args.journal_fit:
        set_paragraph(doc, f"Summary of Scientific Grounds for Consideration at {args.journal}", bold=True, size=11, space_after=6)
        set_paragraph(doc, args.journal_fit, size=11, space_after=12)
    if args.audience:
        set_paragraph(doc, "Summary of Appeal to a Broad Audience", bold=True, size=11, space_after=6)
        set_paragraph(doc, args.audience, size=11, space_after=12)

    set_paragraph(doc, "Items Included in This Submission", bold=True, size=11, space_after=6)
    set_paragraph(doc, f"The manuscript comprises {args.stats}. A Supplementary Information document is included.", size=11, space_after=12)

    if args.declarations:
        set_paragraph(doc, args.declarations, size=11, space_after=12)

    set_paragraph(doc, "We appreciate your time and consideration and look forward to hearing from you.", size=11, space_after=24)
    build_signature(doc, args)


def build_revision(doc, args):
    build_header(doc, args)
    set_paragraph(doc, "Dear Editor,", size=11, space_after=6)

    opening = (f'We are resubmitting our manuscript entitled '
               f'"{args.title}" '
               f'(MS# {args.msid}) to {args.journal}, '
               f'following the reviewers\' comments. '
               f'We have addressed all concerns raised by the reviewers; '
               f'a detailed point-by-point Response Letter accompanies this submission.')
    set_paragraph(doc, opening, size=11, space_after=12)

    if args.changes:
        set_paragraph(doc, "Changes Made", bold=True, size=11, space_after=6)
        set_multiline(doc, args.changes, size=11, space_after=6)
        doc.add_paragraph()

    set_paragraph(doc, "Items Included in This Resubmission", bold=True, size=11, space_after=6)
    set_paragraph(doc, f"The revised manuscript comprises {args.stats}. A Supplementary Information document and a detailed Response Letter are included.", size=11, space_after=12)

    set_paragraph(doc, "We thank the editor and reviewers for their constructive feedback, which has strengthened the manuscript.", size=11, space_after=24)
    build_signature(doc, args)


def build_signature(doc, args):
    set_paragraph(doc, "Sincerely,", size=11, space_after=24)
    set_paragraph(doc, args.authors, size=11, space_after=12)
    if args.corresponding:
        set_paragraph(doc, "* Corresponding authors", size=10, space_after=6)
        set_multiline(doc, args.corresponding, size=10, space_after=6)


def main():
    p = argparse.ArgumentParser(description="Generate a DOCX cover letter")
    p.add_argument("--type", required=True, choices=["first", "revision"])
    p.add_argument("--journal", default="")
    p.add_argument("--title", default="")
    p.add_argument("--out", required=True, help="Output .docx path")

    # First-submission fields
    p.add_argument("--background")
    p.add_argument("--findings")
    p.add_argument("--journal-fit")
    p.add_argument("--audience")
    p.add_argument("--declarations")

    # Revision fields
    p.add_argument("--msid", help="Manuscript ID (revision only)")
    p.add_argument("--changes", help="Non-scientific changes, \\n-separated (revision)")

    # Shared
    p.add_argument("--stats", default="")
    p.add_argument("--authors", default="")
    p.add_argument("--corresponding", default="")
    p.add_argument("--dept", default="")
    p.add_argument("--institution", default="")
    p.add_argument("--address", default="")

    args = p.parse_args()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    if args.type == "first":
        build_first(doc, args)
    else:
        build_revision(doc, args)

    out = args.out
    if not out.endswith(".docx"):
        out += ".docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
